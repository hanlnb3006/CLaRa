#!/usr/bin/env python3
"""Convert a Markdown file to a Word .docx using python-docx.

Supports: ATX headings (#..######), fenced code blocks, GFM tables,
bullet/numbered lists, bold/italic inline, horizontal rules, paragraphs.
Usage: python md_to_docx.py input.md output.docx
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)\*(?!\*)")
INLINE_CODE = re.compile(r"`(.+?)`")


def add_runs(paragraph, text):
    """Add inline-formatted runs to a paragraph."""
    # Tokenize preserving order: code first, then bold, then italic.
    tokens = [text]
    # Process inline code
    new_tokens = []
    for tok in tokens:
        last = 0
        for m in INLINE_CODE.finditer(tok):
            if m.start() > last:
                new_tokens.append(("txt", tok[last:m.start()]))
            new_tokens.append(("code", m.group(1)))
            last = m.end()
        if last < len(tok):
            new_tokens.append(("txt", tok[last:]))
    tokens = new_tokens
    # Process bold within txt tokens
    new_tokens = []
    for kind, tok in tokens:
        if kind != "txt":
            new_tokens.append((kind, tok))
            continue
        last = 0
        for m in INLINE_BOLD.finditer(tok):
            if m.start() > last:
                new_tokens.append(("txt", tok[last:m.start()]))
            new_tokens.append(("bold", m.group(1)))
            last = m.end()
        if last < len(tok):
            new_tokens.append(("txt", tok[last:]))
    tokens = new_tokens
    # Process italic within txt tokens
    new_tokens = []
    for kind, tok in tokens:
        if kind != "txt":
            new_tokens.append((kind, tok))
            continue
        last = 0
        for m in INLINE_ITALIC.finditer(tok):
            if m.start() > last:
                new_tokens.append(("txt", tok[last:m.start()]))
            new_tokens.append(("italic", m.group(1)))
            last = m.end()
        if last < len(tok):
            new_tokens.append(("txt", tok[last:]))
    tokens = new_tokens
    for kind, tok in tokens:
        if not tok:
            continue
        run = paragraph.add_run(tok)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def set_cell_background(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def add_table(doc, rows):
    """rows: list of list of cell strings; first row is header."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, row[j] if j < len(row) else "")
            if i == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_background(cell, "D9E2F3")
    doc.add_paragraph()


def convert(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Remove trailing markdown anchors
            text = re.sub(r"\s*\{.*\}$", "", text).strip()
            doc.add_heading(text, level=min(level, 6))
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            doc.add_paragraph().add_run().add_break()
            i += 1
            continue

        # GFM table
        if "|" in stripped and i + 1 < n and re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", lines[i + 1]):
            table_rows = []
            # header
            header = [c.strip() for c in stripped.strip("|").split("|")]
            table_rows.append(header)
            i += 2  # skip header + separator
            while i < n and "|" in lines[i].strip() and lines[i].strip():
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                table_rows.append(row)
                i += 1
            add_table(doc, table_rows)
            continue

        # Bullet list
        bm = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, bm.group(2))
            i += 1
            continue

        # Numbered list
        nm = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if nm:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, nm.group(2))
            i += 1
            continue

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, stripped.lstrip(">").strip())
            p.runs[0].italic = True if p.runs else None
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Saved {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py input.md output.docx")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])